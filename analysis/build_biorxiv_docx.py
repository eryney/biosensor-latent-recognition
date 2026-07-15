"""Build the bioRxiv submission .docx from the canonical renumbered manuscript.

Source of truth: v3/manuscript_v3.md (renumbered: structure = Figure 1, data
figures = Figures 2-5). Figures are embedded above their caption paragraphs in
the new order. Output: v4/manuscript/manuscript_v4_biorxiv.docx
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]  # repository root
MD = ROOT / "manuscript" / "manuscript_biorxiv.md"
FIG = ROOT / "figures"
OUT = ROOT / "manuscript" / "manuscript_biorxiv.docx"

# Figure-number -> (image file, render width in inches)
FIG_IMAGES = {
    "1": (FIG / "Figure1.png", 6.5),
    "2": (FIG / "fig1_screen_scope_heatmap.png", 6.5),
    "3": (FIG / "fig2_sequence_activity_cliffs.png", 6.5),
    "4": (FIG / "fig3_dose_response_leads.png", 6.3),
    "5": (FIG / "fig4_translational_relevance.png", 6.5),
}

doc = Document()

# ---- base styles: bioRxiv-friendly (11 pt serif body, 1.5 spacing) ----
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(11)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(6)


def _fmt_cite(nums):
    """Collapse a citation number list into ACS superscript text, e.g.
    [2,3,4] -> '2-4', [1] -> '1', [8,4,9,10] -> '4,8-10'."""
    ns = sorted(set(int(n) for n in nums))
    parts, i = [], 0
    while i < len(ns):
        j = i
        while j + 1 < len(ns) and ns[j + 1] == ns[j] + 1:
            j += 1
        if j - i >= 2:
            parts.append(f"{ns[i]}-{ns[j]}")
        else:
            parts.extend(str(ns[k]) for k in range(i, j + 1))
        i = j + 1
    return ",".join(parts)


# Split on **bold** and [CITE:...] markers, keeping the delimiters.
_TOKEN = re.compile(r"(\*\*.*?\*\*|\[CITE:[0-9,]+\])")


def add_runs_with_bold(paragraph, text):
    """Render inline **bold** as bold runs and [CITE:n,m] as superscript numbers."""
    for seg in _TOKEN.split(text):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            r = paragraph.add_run(seg[2:-2])
            r.bold = True
        elif seg.startswith("[CITE:"):
            nums = seg[len("[CITE:"):-1].split(",")
            r = paragraph.add_run(_fmt_cite(nums))
            r.font.superscript = True
        else:
            paragraph.add_run(seg)


def add_figure(num):
    img, width = FIG_IMAGES[num]
    if img.exists():
        doc.add_picture(str(img), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        raise FileNotFoundError(img)


lines = MD.read_text().split("\n")
title_done = False

for raw in lines:
    line = raw.rstrip("\n")
    if not line.strip():
        continue

    # Title (# ...) -> centered bold 16 pt
    if line.startswith("# ") and not title_done:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line[2:].strip())
        r.bold = True
        r.font.size = Pt(16)
        p.paragraph_format.space_after = Pt(12)
        title_done = True
        continue

    # Section heading (## ...)
    if line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=1)
        continue

    # Subsection (### ...)
    if line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=2)
        continue

    # Figure caption paragraph: embed the image first, then the caption text
    m = re.match(r"^\*\*Figure (\d+)\.", line)
    if m:
        num = m.group(1)
        add_figure(num)
        cap = doc.add_paragraph()
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after = Pt(14)
        add_runs_with_bold(cap, line)
        for r in cap.runs:
            r.font.size = Pt(9.5)
        continue

    # Author line / affiliations (right after title, before Abstract): center
    if not any(h in line for h in ("Abstract",)) and title_done and \
       (line.startswith("Eryney") or re.match(r"^\d+ ", line)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line.strip())
        r.font.size = Pt(11) if line.startswith("Eryney") else Pt(10)
        continue

    # Regular paragraph (with inline bold)
    p = doc.add_paragraph()
    add_runs_with_bold(p, line)

doc.save(str(OUT))
print("saved", OUT)
print("paragraphs:", len(doc.paragraphs))
